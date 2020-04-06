import docx
from docx import Document
from docx.shared import Inches
from src.tool.find_all_hdf5 import *
import docx.image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
data_root_path = 'D:/DeepTool/DeepTool/deal-data-whole5/train_procese'
# 创建一个Workbook对象，相当于创建了一个Excel文件
document = Document()
document.add_heading('基于深度学习的熔池正面图像预测研究报告',0)
optimizers = ['sgd','adam','Adagrad','Adadelta','Adamax','Nadam','RMSprop']
loss_fuctions = ['categorical_crossentropy','binary_crossentropy']
models = [ 'VGG16', 'VGG19', 'InceptionV3','Xception', 'InceptionResNetV2', 'DenseNet201']
i = 0
j = 0
for all_txt in all_files(data_root_path, patterns='*epoch.txt'):
    model_fille_name = os.path.basename(all_txt)
    if i == 0:
        temp = model_fille_name
        document.add_heading("1 神经网路" + model_fille_name + "模型训练情况图括", level=0)
    if (temp != model_fille_name) and temp :
        #写入标题
        document.add_heading(str(i+1)+"神经网路"+model_fille_name+"模型训练情况图括",level=0)
        print("模型已经改变")
        print(all_txt)
        i = 0
        #获取父目录
    img_path = os.path.abspath(os.path.join(all_txt, "../"))
    try:
        document.add_picture(img_path + '/filename1.png',Inches(6))
        if i < 14:
            a = document.add_paragraph('图'+str(i+1)+'.'+str(i % 2)+': optimizer is '+str(i%8) +str(optimizers[i%7]) + ' and loss function is '+str(loss_fuctions[i % 2]))
            a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_picture(img_path + '/filename2.png',Inches(6))
        if i < 14:
            b = document.add_paragraph('图'+str(i+1)+'.'+str((i+1) % 2)+': optimizer is '+str(i%8) + str(optimizers[i%7]) + ' and loss function is ' + str(loss_fuctions[i % 2]))
            b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        i = i+1
        print("optimizer is "," ",i)
        temp = model_fille_name
    except docx.image.exceptions.UnrecognizedImageError as e:
        #i = i - 1
        print("异常")
        pass
    except FileNotFoundError:
        print("异常")
        #i = i - 1
        pass
document.save(data_root_path+'基于神单独学习的熔池状态检测模型训练过程图-全.docx')
#for optimizer in optimizers:
 #   for loss_fuction in loss_fuctions:

#D:\DeepTool\DeepTool\deal-data-whole5\train_procese\10000\filename1.png